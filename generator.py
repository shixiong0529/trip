"""
文档生成层
Markdown → HTML (Jinja2) → PDF (WeasyPrint) / DOCX (python-docx)

PDF 生成依赖 WeasyPrint，需要系统级依赖（pango、cairo、gobject）。
macOS 上可通过 `brew install pango cairo glib` 安装。
如果环境不支持，PDF 功能将不可用，但 HTML 和 DOCX 不受影响。
"""

import re
import io
import logging
from decimal import Decimal, InvalidOperation
from pathlib import Path

from jinja2 import Environment, FileSystemLoader

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from services.route_map import normalize_route_overview

logger = logging.getLogger(__name__)


_SECTION_ORDER = (
    "weather", "transport", "lodging", "itinerary", "budget",
    "booking", "pitfalls", "outbound", "packing", "knowledge",
)
_REQUIRED_SECTIONS = set(_SECTION_ORDER) - {"outbound"}
_SECTION_TITLES = {
    "weather": "🌤️ 天气与穿搭",
    "transport": "🚄 城际交通建议",
    "lodging": "🏨 住宿推荐",
    "itinerary": "📅 分日行程",
    "budget": "💰 总预算拆解",
    "booking": "🚨 必做预约 & 证件清单",
    "pitfalls": "⚠️ 避坑提示",
    "outbound": "🌍 出境须知",
    "packing": "🎒 行前物品清单",
    "knowledge": "🌳 行程知识图谱",
}
_TABLE_SCHEMAS = {
    "transport": (
        ("方向", r"方向|行程"),
        ("推荐方式", r"方式|交通|工具"),
        ("耗时", r"耗时|时长|时间"),
        ("参考价（人均）", r"参考价|价格|费用|票价"),
        ("提示", r"提示|建议|备注"),
    ),
    "lodging": (
        ("档位", r"档位|类型|级别"),
        ("推荐区域/酒店", r"区域|酒店|住宿|推荐"),
        ("区位优势", r"区位|优势|位置|特色"),
        ("人均/晚参考价", r"人均|参考价|价格|房价"),
        ("适合人群", r"适合|人群"),
    ),
    "itinerary": (
        ("时段", r"时段|时间"),
        ("安排", r"安排|活动|行程|项目"),
        ("耗时", r"耗时|时长"),
        ("提示", r"提示|建议|备注|费用"),
    ),
    "daily_budget": (
        ("项目", r"项目|类别"),
        ("明细", r"明细|说明|计算|内容"),
        ("费用", r"费用|金额|小计|合计"),
        ("备注", r"备注|提示"),
    ),
    "budget": (
        ("类别", r"类别|项目|分类"),
        ("明细", r"明细|说明|计算"),
        ("人均", r"人均"),
        ("合计", r"合计|总价|总计"),
    ),
    "outbound": (
        ("项目", r"项目|类别"),
        ("详情", r"详情|说明|建议"),
    ),
    "packing": (
        ("类别", r"类别|分类"),
        ("物品（写出具体数量）", r"物品|清单|数量|内容"),
    ),
}


_MONEY_RANGE_RE = re.compile(
    r"[¥￥]\s*(?P<low>\d[\d,]*(?:\.\d+)?)"
    r"(?:\s*(?:-|–|—|~|至)\s*[¥￥]?\s*(?P<high>\d[\d,]*(?:\.\d+)?))?"
)
_DAILY_BUDGET_TOTAL_RE = re.compile(
    r"(?P<prefix>本日预算.{0,80}?[¥￥]\s*)"
    r"(?P<amount>\d[\d,]*(?:\.\d+)?"
    r"(?:\s*(?:-|–|—|~|至)\s*[¥￥]?\s*\d[\d,]*(?:\.\d+)?)?)"
)


def correct_daily_budget_totals(markdown_content: str) -> str:
    """用明细复算每个“本日预算”的合计，仅替换错误的金额文本。

    模型擅长生成行程，但并不可靠地执行算术。这里把括号内以 ``+``、
    ``、`` 或分号分隔的费用视为权威明细，支持单价乘人数、显式等式和
    金额区间。无法可靠识别至少两项明细时保持原文，避免猜测。
    """
    return "\n".join(_correct_daily_budget_line(line) for line in markdown_content.split("\n"))


def normalize_report_markdown(markdown_content: str) -> str:
    """统一模型输出中会影响文档格式的结构，同时保留原始内容。"""
    markdown_content = correct_daily_budget_totals(markdown_content)
    markdown_content = _canonicalize_headings(markdown_content)
    markdown_content = _normalize_and_order_sections(markdown_content)
    section_re = re.compile(
        r"(?P<heading>^##[^\n]*行程知识图谱[^\n]*\n)"
        r"(?P<body>.*?)(?=^##\s|^---\s*$|^>\s|\Z)",
        re.MULTILINE | re.DOTALL,
    )

    def normalize_section(match: re.Match) -> str:
        heading = match.group("heading")
        body = match.group("body").strip()
        fence = re.search(r"```[^\n]*\n?(.*?)```", body, re.DOTALL)
        if fence:
            graph = _normalize_knowledge_graph(fence.group(1))
            body = body[:fence.start()] + f"```\n{graph}\n```" + body[fence.end():]
        elif re.search(r"\bDay\s*\d+\b", body, re.IGNORECASE):
            body = f"```\n{_normalize_knowledge_graph(body)}\n```"
        return heading + body.strip() + "\n\n"

    return section_re.sub(normalize_section, markdown_content)


def _canonicalize_headings(markdown_content: str) -> str:
    lines = []
    for line in markdown_content.splitlines():
        heading = re.match(r"^(#{2,4})\s+(.+?)\s*$", line)
        if heading:
            level, title = heading.groups()
            day = re.match(
                r"^(?:Day\s*|第\s*)(\d+)(?:\s*[天日])?\s*[-—:：·.]?\s*(.*)$",
                re.sub(r"[*_]+", "", title).strip(),
                re.IGNORECASE,
            )
            if day:
                theme = day.group(2).strip(" -—:：·.")
                lines.append(f"### Day {int(day.group(1))}" + (f" · {theme}" if theme else ""))
                continue

            key = _section_key(title)
            # H3/H4 中“预约清单”等是合法子标题，只提升明确的顶级板块名。
            if key and (level == "##" or _is_explicit_top_level_title(title, key)):
                lines.append("## " + _SECTION_TITLES[key])
                continue
        lines.append(line)
    suffix = "\n" if markdown_content.endswith("\n") else ""
    return "\n".join(lines) + suffix


def _section_key(title: str) -> str | None:
    compact = re.sub(r"[\s*_#·&🌤️🚄🏨📅💰🚨⚠️🌍🎒🌳]+", "", title)
    if "知识图谱" in compact:
        return "knowledge"
    if "分日行程" in compact or "详细行程" in compact or "每日行程" in compact:
        return "itinerary"
    if "总预算" in compact or "预算拆解" in compact or "费用汇总" in compact:
        return "budget"
    if "天气" in compact or "穿搭" in compact:
        return "weather"
    if "住宿" in compact or "酒店推荐" in compact:
        return "lodging"
    if "交通" in compact and "市内" not in compact:
        return "transport"
    if "必做预约" in compact or ("预约" in compact and "证件" in compact):
        return "booking"
    if "避坑" in compact or "风险提示" in compact:
        return "pitfalls"
    if "出境须知" in compact or "出境指南" in compact:
        return "outbound"
    if ("物品" in compact or "行李" in compact) and "清单" in compact:
        return "packing"
    return None


def _is_explicit_top_level_title(title: str, key: str) -> bool:
    compact = re.sub(r"[\s*_#·&🌤️🚄🏨📅💰🚨⚠️🌍🎒🌳]+", "", title)
    explicit = {
        "weather": ("天气与穿搭",),
        "transport": ("城际交通建议",),
        "lodging": ("住宿推荐",),
        "itinerary": ("分日行程", "详细行程", "每日行程"),
        "budget": ("总预算拆解", "费用汇总"),
        "booking": ("必做预约证件清单",),
        "pitfalls": ("避坑提示", "风险提示"),
        "outbound": ("出境须知", "出境指南"),
        "packing": ("行前物品清单", "行李清单"),
        "knowledge": ("行程知识图谱",),
    }
    return any(name in compact for name in explicit[key])


def _normalize_and_order_sections(markdown_content: str) -> str:
    section_matches = list(re.finditer(r"^##\s+(.+?)\s*$", markdown_content, re.MULTILINE))
    if not section_matches:
        return markdown_content

    preamble = markdown_content[:section_matches[0].start()].rstrip()
    known: dict[str, list[str]] = {}
    unknown: list[str] = []
    recognized_keys = []
    for index, match in enumerate(section_matches):
        end = section_matches[index + 1].start() if index + 1 < len(section_matches) else len(markdown_content)
        title = match.group(1).strip()
        body = markdown_content[match.end():end].strip()
        key = _section_key(title)
        if key:
            recognized_keys.append(key)
            body = _normalize_section_tables(body, key)
            known.setdefault(key, []).append(body)
        else:
            unknown.append(f"## {title}\n\n{body}".strip())

    # 小型 Markdown 片段保持原样；完整攻略才启用固定九板块契约。
    if len(set(recognized_keys)) < 5:
        return markdown_content

    parts = [preamble] if preamble else []
    for key in _SECTION_ORDER:
        bodies = known.get(key, [])
        if not bodies and key not in _REQUIRED_SECTIONS:
            continue
        body = "\n\n".join(part for part in bodies if part).strip()
        if not body:
            body = "> 本板块未生成有效内容，请以官方实时信息为准。"
        parts.append(f"## {_SECTION_TITLES[key]}\n\n{body}")
        if key == "packing" and unknown:
            parts.extend(unknown)
            unknown = []
    parts.extend(unknown)
    return "\n\n".join(parts).rstrip() + "\n"


def _normalize_section_tables(body: str, section_key: str) -> str:
    default_schema = _TABLE_SCHEMAS.get(section_key)
    if not default_schema:
        return body

    # 只允许水平空白，避免 ``\s*`` 吞掉表格后的空行并制造一行空数据。
    table_re = re.compile(r"(?:^[ \t]*\|.*\|[ \t]*$\n?)+", re.MULTILINE)

    def normalize_table(match: re.Match) -> str:
        rows = _parse_table(match.group(0).splitlines())
        if not rows:
            return match.group(0)
        source_headers = [re.sub(r"[*_`]+", "", cell).strip() for cell in rows[0]]
        schema = default_schema
        if section_key == "itinerary":
            header_text = " ".join(source_headers)
            if not re.search(r"时段|时间", header_text) and re.search(
                r"项目|类别|费用|金额|明细|合计", header_text
            ):
                schema = _TABLE_SCHEMAS["daily_budget"]
        used: set[int] = set()
        mapping: list[int | None] = []
        # 先完成所有语义匹配，再做位置兜底；否则三列表中的“建议”可能被
        # 前一个缺失的“耗时”抢占，导致内容落到错误列。
        for _, aliases in schema:
            source_index = next(
                (idx for idx, header in enumerate(source_headers) if idx not in used and re.search(aliases, header)),
                None,
            )
            if source_index is not None:
                used.add(source_index)
            mapping.append(source_index)
        for position, source_index in enumerate(mapping):
            if source_index is not None:
                continue
            fallback = position if position < len(source_headers) and position not in used else next(
                (idx for idx in range(len(source_headers)) if idx not in used), None
            )
            if fallback is not None:
                mapping[position] = fallback
                used.add(fallback)

        normalized_rows = []
        for source_row in rows[1:]:
            row = [source_row[idx] if idx is not None and idx < len(source_row) else "" for idx in mapping]
            extras = [source_row[idx] for idx in range(len(source_row)) if idx not in used and source_row[idx]]
            if extras:
                row[-1] = "；".join(filter(None, [row[-1], *extras]))
            normalized_rows.append(row)

        headers = [header for header, _ in schema]
        separator = ["---"] * len(headers)
        all_rows = [headers, separator, *normalized_rows]
        return "\n".join("| " + " | ".join(row) + " |" for row in all_rows) + "\n"

    return table_re.sub(normalize_table, body).rstrip()


def _normalize_knowledge_graph(content: str) -> str:
    """将知识图谱稳定为根节点、逐日分支和汇总行组成的等宽文本树。"""
    text = re.sub(r"[ \t]*\n[ \t]*", " ", content.strip())
    if not re.search(r"\bDay\s*\d+\b", text, re.IGNORECASE):
        return content.strip()

    # 兼容模型输出的 ├──、|——、│──，以及完全漏掉树枝符号的 Day N。
    branch_marker = r"(?:[|│├└]\s*)?[─—-]{2,}"
    text = re.sub(
        rf"\s*{branch_marker}\s*(?=Day\s*\d+\b)",
        "\n@@DAY@@",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        rf"\s*{branch_marker}\s*(?=🏁)", "\n@@FINISH@@", text
    )
    text = re.sub(
        r"(?<!@@DAY@@)(?<!^)[ \t]+(?=Day\s*\d+\s*[·.])",
        "\n@@DAY@@",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\s+(?=[💰📏🏔️👥✅])", "\n@@SUMMARY@@", text)

    root_parts: list[str] = []
    day_parts: list[str] = []
    finish_parts: list[str] = []
    summary_parts: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("@@DAY@@"):
            day_parts.append(line.removeprefix("@@DAY@@").strip())
        elif line.startswith("@@FINISH@@"):
            finish_parts.append(line.removeprefix("@@FINISH@@").strip())
        elif line.startswith("@@SUMMARY@@"):
            summary_parts.append(line.removeprefix("@@SUMMARY@@").strip())
        elif re.match(r"Day\s*\d+\b", line, re.IGNORECASE):
            day_parts.append(line)
        elif line[0] in "💰📏🏔️👥✅":
            summary_parts.append(line)
        else:
            root_parts.append(line)

    if not day_parts:
        return content.strip()

    result = [" ".join(root_parts).strip()] if root_parts else []
    has_finish = bool(finish_parts)
    for index, day in enumerate(day_parts):
        is_last_without_finish = index == len(day_parts) - 1 and not has_finish
        result.append(("└── " if is_last_without_finish else "├── ") + day)
    result.extend("└── " + finish for finish in finish_parts)
    if summary_parts:
        result.append("")
        result.extend(summary_parts)
    return "\n".join(result)


def _correct_daily_budget_line(line: str) -> str:
    if "本日预算" not in line:
        return line

    declared = _DAILY_BUDGET_TOTAL_RE.search(line)
    if not declared:
        return line

    tail = line[declared.end():]
    open_positions = [pos for pos in (tail.find("（"), tail.find("(")) if pos >= 0]
    if not open_positions:
        return line
    open_pos = min(open_positions)
    opener = tail[open_pos]
    closer = "）" if opener == "（" else ")"
    close_pos = tail.find(closer, open_pos + 1)
    if close_pos < 0:
        return line

    detail = tail[open_pos + 1:close_pos]
    components = re.split(r"\s*(?:\+|＋|；|;|、|，|\s和\s)\s*", detail)
    costs = [cost for part in components if (cost := _component_cost(part)) is not None]
    if len(costs) < 2:
        return line

    low_total = sum((cost[0] for cost in costs), Decimal("0"))
    high_total = sum((cost[1] for cost in costs), Decimal("0"))
    corrected = _format_money(low_total)
    if high_total != low_total:
        corrected += "-¥" + _format_money(high_total)

    start, end = declared.span("amount")
    return line[:start] + corrected + line[end:]


def _component_cost(component: str) -> tuple[Decimal, Decimal] | None:
    component = component.strip()
    if not component or re.search(r"人均|合计|总计", component):
        return None

    # “单价×人数=合计”以等号右侧结果为准，避免单价与结果重复计入。
    equation_parts = re.split(r"[=＝]", component)
    if len(equation_parts) > 1:
        return _first_money_range(equation_parts[-1])

    money_match = _MONEY_RANGE_RE.search(component)
    if not money_match:
        return None
    money = _money_match_to_range(money_match)
    if money is None:
        return None

    before = component[:money_match.start()]
    after = component[money_match.end():]
    quantity_before = re.search(
        r"(\d+(?:\.\d+)?)\s*(?:人|晚|天|张|间|辆|次|份|套|位)?\s*[×xX*]\s*$",
        before,
    )
    quantity_after = re.match(
        r"\s*(?:/\s*(?:人|晚|天|张|间|辆|次|份|套|位))?"
        r"\s*[×xX*]\s*(\d+(?:\.\d+)?)",
        after,
    )
    quantity_text = (
        quantity_before.group(1) if quantity_before else
        quantity_after.group(1) if quantity_after else None
    )
    if quantity_text:
        try:
            quantity = Decimal(quantity_text)
            return money[0] * quantity, money[1] * quantity
        except InvalidOperation:
            return None
    return money


def _first_money_range(text: str) -> tuple[Decimal, Decimal] | None:
    match = _MONEY_RANGE_RE.search(text)
    return _money_match_to_range(match) if match else None


def _money_match_to_range(match: re.Match) -> tuple[Decimal, Decimal] | None:
    try:
        low = Decimal(match.group("low").replace(",", ""))
        high_text = match.group("high")
        high = Decimal(high_text.replace(",", "")) if high_text else low
        return low, high
    except (InvalidOperation, AttributeError):
        return None


def _format_money(value: Decimal) -> str:
    if value == value.to_integral_value():
        return f"{int(value):,}"
    return f"{value:,.2f}".rstrip("0").rstrip(".")

# WeasyPrint 延迟导入，环境不支持时优雅降级
_weasyprint_available = False
try:
    from weasyprint import HTML as WeasyHTML
    _weasyprint_available = True
except (ImportError, OSError) as e:
    logger.warning(f"WeasyPrint 不可用（PDF 功能将受限）: {e}")


# ---------- Markdown 解析器 ----------
def _parse_markdown_to_blocks(md: str) -> list[dict]:
    """将 Markdown 文本解析为结构化块"""
    blocks = []
    lines = md.split("\n")
    i = 0

    # 跳过开头空行
    while i < len(lines) and not lines[i].strip():
        i += 1

    while i < len(lines):
        line = lines[i]

        # 一级标题 #
        m = re.match(r"^#\s+(.+)", line)
        if m:
            blocks.append({"type": "h1", "content": m.group(1).strip()})
            i += 1
            continue

        # 二级标题 ##
        m = re.match(r"^##\s+(.+)", line)
        if m:
            blocks.append({"type": "h2", "content": m.group(1).strip()})
            i += 1
            continue

        # 三级标题 ###
        m = re.match(r"^###\s+(.+)", line)
        if m:
            blocks.append({"type": "h3", "content": m.group(1).strip()})
            i += 1
            continue

        # 引用 >
        if line.startswith("> "):
            ref_lines = []
            while i < len(lines) and lines[i].startswith("> "):
                ref_lines.append(lines[i][2:])
                i += 1
            blocks.append({"type": "blockquote", "content": "\n".join(ref_lines)})
            continue

        # 代码块 ```
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            blocks.append({"type": "code", "content": "\n".join(code_lines)})
            continue

        # 表格（检测 | 开头）
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "table", "rows": _parse_table(table_lines)})
            continue

        # 任务列表 - [ ] xxx 或 - [x] xxx（必须先于无序列表判断，否则会被当成普通列表）
        if re.match(r"^[\s]*[-*]\s+\[[ x]\]", line):
            task_items = []
            while i < len(lines) and re.match(r"^[\s]*[-*]\s+\[[ x]\]", lines[i]):
                m = re.match(r"^[\s]*[-*]\s+\[([ x])\]\s*(.*)", lines[i])
                if m:
                    task_items.append({"checked": m.group(1) == "x", "text": m.group(2)})
                i += 1
            blocks.append({"type": "tasks", "items": task_items})
            continue

        # 无序列表 - xxx 或 * xxx（支持两级嵌套：缩进 >=2 空格的子项挂到上一个顶层项下）
        if re.match(r"^[\s]*[-*]\s+", line):
            list_items = []
            while i < len(lines) and re.match(r"^[\s]*[-*]\s+", lines[i]) \
                    and not re.match(r"^[\s]*[-*]\s+\[[ x]\]", lines[i]):
                m_item = re.match(r"^(\s*)[-*]\s+(.*)", lines[i])
                # tab 缩进按 2 空格折算，同样识别为子项
                indent = len(m_item.group(1).replace("\t", "  "))
                item_text = m_item.group(2)
                if indent >= 2 and list_items:
                    list_items[-1]["children"].append(item_text)
                else:
                    list_items.append({"text": item_text, "children": []})
                i += 1
            blocks.append({"type": "ul", "items": list_items})
            continue

        # 有序列表 1. xxx
        if re.match(r"^[\s]*\d+\.\s+", line):
            list_items = []
            while i < len(lines) and re.match(r"^[\s]*\d+\.\s+", lines[i]):
                item_text = re.sub(r"^[\s]*\d+\.\s+", "", lines[i])
                list_items.append(item_text)
                i += 1
            blocks.append({"type": "ol", "items": list_items})
            continue

        # 水平线 ---
        if line.strip() in ("---", "***", "___"):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", ">", "```", "|", "- ", "* ", "1. ")):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"type": "p", "content": "\n".join(para_lines)})

    return blocks


def _parse_table(lines: list[str]) -> list[list[str]]:
    """解析 Markdown 表格"""
    rows = []
    for line in lines:
        # 跳过分隔行 |---|
        if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if rows:
        width = max(len(row) for row in rows)
        while len(rows[0]) < width:
            rows[0].append(f"详情{len(rows[0]) + 1}")
        for row in rows:
            row.extend([""] * (width - len(row)))
    return rows


# ---------- HTML 生成 ----------
def _blocks_to_html_fragment(blocks: list[dict]) -> str:
    """将块转为 HTML Report Generator 风格的片段"""
    html = ""
    i = 0
    in_section = False  # 追踪是否在 section div 内
    in_container = False  # 追踪 hero 后的 container div 是否已打开

    while i < len(blocks):
        block = blocks[i]
        t = block["type"]
        content = block.get("content", "")

        # H1 标题 → Hero 区
        if t == "h1":
            html += _render_hero(content)
            in_container = True
            # H1 后到第一个章节之间是概览区。模型偶尔会把加粗标记拆成
            # 独立段落，因此不限制段落数量，再由 _render_overview 过滤。
            overview_paras = []
            j = i + 1
            while j < len(blocks) and blocks[j]["type"] == "p":
                overview_paras.append(blocks[j]["content"])
                j += 1
            if overview_paras:
                html += _render_overview(overview_paras)
                i = j - 1

        # H2 分节标题 → Section
        elif t == "h2":
            if _is_route_map_section(content):
                i += 1
                while i < len(blocks) and blocks[i]["type"] != "h2":
                    i += 1
                continue
            if in_section:
                html += "</div>\n"  # 关闭上一个 section
            slug = re.sub(r"[^\w\u4e00-\u9fff]+", "-", content).strip("-")
            # \u6807\u9898\u81ea\u5e26 emoji \u65f6\u4e0d\u518d\u8ffd\u52a0\u56fe\u6807\uff0c\u907f\u514d\u53cc\u56fe\u6807
            if re.match(r"^[\w\u4e00-\u9fff]", content):
                icon_html = f'<span class="icon">{_section_icon(content)}</span>'
            else:
                icon_html = ""
            html += f'<div class="section" id="{slug}">\n'
            html += f'<h2 class="section-title">{icon_html}{_inline_md(content)}</h2>\n'
            in_section = True

        # H3 "Day X" → 日程卡片
        elif t == "h3" and ("Day" in content or content.startswith("Day ") or "day" in content.lower()):
            day_html, new_i = _render_day_card(blocks, i)
            html += day_html
            i = new_i
            continue

        # 其他 H3（预约清单、证件等小节标题）
        elif t == "h3":
            html += f'<h3 class="sub-title">{_inline_md(content)}</h3>\n'

        # 表格
        elif t == "table":
            html += _render_table(block["rows"])

        # 列表（支持两级嵌套）
        elif t == "ul":
            items = ""
            for it in block["items"]:
                items += f"<li>{_inline_md(it['text'])}"
                children = it.get("children") or []
                if children:
                    child_html = "".join(f"<li>{_inline_md(c)}</li>" for c in children)
                    items += f"<ul>{child_html}</ul>"
                items += "</li>"
            html += f"<div class=\"card\"><ul>{items}</ul></div>\n"

        elif t == "ol":
            items = "".join(f"<li>{_inline_md(it)}</li>" for it in block["items"])
            html += f"<div class=\"card\"><ol>{items}</ol></div>\n"

        elif t == "tasks":
            items = ""
            for it in block["items"]:
                chk = 'checked disabled' if it["checked"] else ""
                items += f'<li class="task-item"><input type="checkbox" {chk}/><span>{_inline_md(it["text"])}</span></li>'
            html += f'<div class="card"><ul class="task-list">{items}</ul></div>\n'

        # 代码块（通常是知识图谱树）
        elif t == "code":
            html += f'<div class="tree">{_escape_html(_normalize_knowledge_graph(content))}</div>\n'

        # 引用块
        elif t == "blockquote":
            html += f'<blockquote>{_inline_md(content)}</blockquote>\n'
            if "免责声明" in content or "免责申明" in content:
                break

        # 普通段落
        elif t == "p":
            # 重要提示 / 预算说明 等段落用 card 包裹
            html += _render_paragraph_card(content)

        i += 1

    if in_section:
        html += "</div>\n"  # 关闭最后一个 section
    if in_container:
        html += "</div>\n"  # 关闭 hero 后打开的 container
    return html


def _render_hero(title: str) -> str:
    """渲染 Hero 区"""
    return f'''<div class="container">
<div class="hero">
  <h1>{_inline_md(title)}</h1>
  <p class="meta">AI 旅行管家 · 路小仙（Leo）定制</p>
</div>
'''


def _is_route_map_section(title: str) -> bool:
    """识别并丢弃模型偶尔输出的路线图章节。"""
    compact = re.sub(r"\s+", "", title)
    return "路线图" in compact or "线路示意图" in compact


def _render_overview(paras: list[str]) -> str:
    """渲染概览区：只保留路线总览 + 重要提示。

    模型不一定会在三段概览之间输出空行。解析器会把相邻行合并成一个
    paragraph，也可能在 ``**`` 和标题之间插入空格。这里按语义标签再次
    切分并过滤孤立 Markdown 标记，避免开头出现空白卡片或裸星号。
    """
    segments: list[str] = []
    overview_label = re.compile(
        r"(?=(?:\*\*\s*)?(?:路线总览|路线纵览|线路(?:总览|纵览))"
        r"|(?:\*\*\s*)?重要提示)"
    )
    for paragraph in paras:
        segments.extend(
            part.strip()
            for part in overview_label.split(paragraph)
            if part.strip() and not re.fullmatch(r"[*_\s]+", part)
        )

    route = next(
        (part for part in segments if re.search(r"路线(?:总览|纵览)|线路(?:总览|纵览)", part)),
        None,
    )
    important = next((part for part in segments if "重要提示" in part), None)

    # 兼容极简 Markdown：完全没有规范概览标签时，保留普通开场说明。
    if not route and not important:
        return "".join(
            _render_paragraph_card(part)
            for part in segments
            if len(_extract_stats(part)) < 2
        )

    html = _render_route_overview_card(route) if route else ""
    html += _render_important_note_card(important) if important else ""
    return html


def _render_paragraph_card(content: str) -> str:
    """根据段落语义渲染重点卡片。"""
    if re.search(r"路线(?:总览|纵览)|线路(?:总览|纵览)", content):
        return _render_route_overview_card(content)
    if "重要提示" in content:
        return _render_important_note_card(content)
    return f'<div class="card"><p>{_inline_md(content)}</p></div>\n'


def _render_route_overview_card(content: str) -> str:
    normalized = normalize_route_overview(content)
    return f'<div class="route-overview-card"><p>{_inline_md(normalized)}</p></div>\n'


def _render_important_note_card(content: str) -> str:
    """规范重要提示标签，清除模型输出中错位的 Markdown 星号。"""
    detail = re.sub(
        r"^\s*\*\*?\s*重要提示\s*\*\*?\s*[:：]?\s*", "", content
    )
    detail = re.sub(r"^\s*重要提示\s*\*{0,2}\s*[:：]?\s*", "", detail)
    detail = detail.strip().strip("*").strip()
    label = '<strong>重要提示：</strong>'
    text = f"{label}{_inline_md(detail)}" if detail else label
    return f'<div class="important-note-card"><p>{text}</p></div>\n'


def _extract_stats(text: str) -> list[tuple[str, str]]:
    """从概览段落中提取关键统计"""
    stats = []
    # 天数
    m = re.search(r"(\d+)\s*天", text)
    if m:
        stats.append(("旅行天数", m.group(1)))
    # 里程
    m = re.search(r"([\d,]+)\s*(km|公里)", text, re.I)
    if m:
        stats.append(("总里程", m.group(1) + "km"))
    # 人数
    m = re.search(r"(\d+)\s*人", text)
    if m:
        stats.append(("出行人数", m.group(1) + "人"))
    # 海拔
    m = re.search(r"(\d+,?\d+)\s*m", text)
    if m:
        stats.append(("最高海拔", m.group(1).replace(",", "") + "m"))
    # 预算（兼容"人均预算 ¥1,850"、"人均约¥2000"等写法）
    m = re.search(r"人均(?:预算)?[约\s]*[¥￥]?\s*([\d,]+)", text)
    if m:
        stats.append(("人均预算", "¥" + m.group(1)))
    # 核心景点数
    m = re.search(r"核心景点[约\s]*(\d+)", text)
    if m:
        stats.append(("核心景点", m.group(1) + "个"))
    return stats


def _section_icon(title: str) -> str:
    """根据标题返回 emoji 图标"""
    mapping = {
        "天气": "🌤️",
        "交通": "🚄",
        "住宿": "🏨",
        "分日": "📅",
        "预算": "💰",
        "预约": "🚨",
        "证件": "📄",
        "避坑": "⚠️",
        "物品": "🎒",
        "知识": "🌳",
        "出境": "🌍",
    }
    for key, icon in mapping.items():
        if key in title:
            return icon
    return "📌"


def _render_day_card(blocks: list[dict], start_idx: int) -> tuple[str, int]:
    """渲染一个 Day 卡片，返回 HTML 和新的索引"""
    title = _inline_md(blocks[start_idx].get("content", ""))

    # 分离标题中的路线和里程信息
    # 例如：Day 1 · 高速启程 永州 → 成都 · 1,150km · 约 13h
    route_info = ""
    main_title = title
    if "·" in title:
        parts = title.split("·")
        main_title = parts[0].strip()
        route_info = " · ".join(p.strip() for p in parts[1:])

    html = '<div class="day-card">\n'
    html += f'  <div class="day-header">\n'
    html += f'    <h3>{main_title}</h3>\n'
    if route_info:
        html += f'    <div class="day-route">{route_info}</div>\n'
    html += f'  </div>\n'

    # 查找紧随其后的表格（时段表）
    i = start_idx + 1
    if i < len(blocks) and blocks[i]["type"] == "table":
        html += '  <div class="day-body">\n'
        html += _render_table(blocks[i]["rows"])
        html += '  </div>\n'
        i += 1

    # 收集本日亮点 / 本日预算作为 footer
    # 注意：两行相邻时会被解析为同一个段落块，需按行拆分成独立徽章
    footer_items = []
    while i < len(blocks) and blocks[i]["type"] == "p":
        p = blocks[i]["content"]
        if not (p.startswith("🎯") or p.startswith("💰")):
            break
        for line in p.split("\n"):
            line = line.strip()
            if line.startswith("🎯") or line.startswith("💰"):
                badge_class = "badge-green" if line.startswith("🎯") else "badge-orange"
                # 移除 emoji
                text = line[1:].strip()
                footer_items.append(f'<span class="badge {badge_class}">{_inline_md(text)}</span>')
            elif line:
                footer_items.append(f'<span class="badge badge-blue">{_inline_md(line)}</span>')
        i += 1

    if footer_items:
        html += '  <div class="day-footer">\n    ' + "\n    ".join(footer_items) + "\n  </div>\n"

    html += "</div>\n"
    return html, i  # i 指向下一个未处理的块，跳过所有已消费内容


def _render_table(rows: list[list[str]]) -> str:
    """渲染表格"""
    if not rows:
        return ""
    headers = rows[0]
    th = "<tr>" + "".join(f"<th>{_inline_md(c)}</th>" for c in headers) + "</tr>"
    trs = ""
    for row in rows[1:]:
        cells = []
        for index, cell in enumerate(row):
            label = _escape_html(_strip_inline_md(headers[index])) if index < len(headers) else ""
            cells.append(f'<td data-label="{label}">{_inline_md(cell)}</td>')
        trs += "<tr>" + "".join(cells) + "</tr>"
    wrapper_class = "table-wrapper"
    if len(rows[0]) == 2:
        wrapper_class += " kv-table-wrapper"
    return f'<div class="{wrapper_class}"><table><thead>{th}</thead><tbody>{trs}</tbody></table></div>\n'



def _escape_html(text: str) -> str:
    """HTML 特殊字符转义，防止 LLM 输出破坏页面结构或注入脚本"""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _inline_md(text: str) -> str:
    """行内 Markdown 转 HTML（先转义，再应用加粗、斜体、代码）"""
    text = _escape_html(text)
    # 加粗 **text**
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    # 斜体 *text*
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<em>\1</em>", text)
    # 行内代码 `code`
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    return text


def _strip_inline_md(text: str) -> str:
    """去掉行内 Markdown 标记（DOCX 纯文本单元格用）"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _add_md_runs(p, text: str):
    """将带 **加粗** 标记的文本按 run 添加到段落，加粗生效而非输出星号"""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(_strip_inline_md(part[2:-2]))
            run.bold = True
        elif part:
            p.add_run(_strip_inline_md(part))


# ---------- 文档生成器 ----------
class TravelGuideGenerator:
    """旅游攻略文档生成器"""

    def __init__(self, templates_dir: str):
        self.templates_dir = templates_dir
        self.env = Environment(
            loader=FileSystemLoader(templates_dir),
            autoescape=True,
        )

    def to_html(self, markdown_content: str, guide_id: str) -> str:
        """Markdown → 完整 HTML 页面"""
        markdown_content = normalize_report_markdown(markdown_content)
        blocks = _parse_markdown_to_blocks(markdown_content)
        body_html = _blocks_to_html_fragment(blocks)

        # 尝试从 h1 提取标题
        title = "旅行攻略"
        for block in blocks:
            if block["type"] == "h1":
                title = block["content"]
                break

        # 渲染 Jinja2 模板
        try:
            template = self.env.get_template("guide.html")
            return template.render(
                guide_id=guide_id,
                title=title,
                body_html=body_html,
            )
        except Exception:
            # 模板不存在时，使用内置模板
            return self._builtin_html(title, body_html, guide_id)

    def to_pdf(self, html_content: str, guide_id: str) -> bytes:
        """HTML → PDF，需要 WeasyPrint 系统依赖"""
        if not _weasyprint_available:
            raise RuntimeError(
                "PDF 生成需要 WeasyPrint。"
                "macOS: brew install pango cairo glib\n"
                "Linux: apt install libpango-1.0-0 libcairo2 libgobject-2.0-0"
            )
        doc = WeasyHTML(string=html_content)
        return doc.write_pdf(
            presentational_hints=True,
        )

    def to_docx(self, markdown_content: str, guide_id: str) -> bytes:
        """Markdown → DOCX"""
        markdown_content = normalize_report_markdown(markdown_content)
        blocks = _parse_markdown_to_blocks(markdown_content)
        doc = Document()

        # 页面设置
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        for block in blocks:
            self._add_docx_block(doc, block)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def _add_docx_block(self, doc: Document, block: dict):
        """将块添加到 Word 文档"""
        t = block["type"]
        text = block.get("content", "").strip()

        if t == "h1":
            p = doc.add_heading(_strip_inline_md(text), level=1)
        elif t == "h2":
            p = doc.add_heading(_strip_inline_md(text), level=2)
        elif t == "h3":
            p = doc.add_heading(_strip_inline_md(text), level=3)
        elif t == "p":
            p = doc.add_paragraph()
            _add_md_runs(p, text)
        elif t == "blockquote":
            p = doc.add_paragraph(_strip_inline_md(text))
            p.paragraph_format.left_indent = Cm(1)
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 116, 139)
        elif t == "ul":
            for item in block.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                _add_md_runs(p, item["text"])
                for child in item.get("children") or []:
                    cp = doc.add_paragraph(style="List Bullet 2")
                    _add_md_runs(cp, child)
        elif t == "ol":
            for item in block.get("items", []):
                p = doc.add_paragraph(style="List Number")
                _add_md_runs(p, item)
        elif t == "tasks":
            for item in block.get("items", []):
                prefix = "☑ " if item["checked"] else "☐ "
                p = doc.add_paragraph(prefix)
                _add_md_runs(p, item["text"])
        elif t == "code":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif t == "table":
            rows = block.get("rows", [])
            if not rows:
                return
            # LLM 输出的表格可能行宽不一致，按最宽行建表，缺的补空
            n_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=n_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci in range(n_cols):
                    cell_text = row[ci] if ci < len(row) else ""
                    cell = table.cell(ri, ci)
                    cell.text = _strip_inline_md(cell_text)
                    if ri == 0:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True
        elif t == "hr":
            doc.add_paragraph("─" * 40)

    def _builtin_html(self, title: str, body_html: str, guide_id: str) -> str:
        """内置 HTML 模板（当 Jinja2 模板文件不可用时）"""
        return f"""<!DOCTYPE html>
<html lang="zh">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{_escape_html(title)}</title>
<style>
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif; background:#f8f9fa; color:#1a1a2e; line-height:1.8; }}
  .guide-container {{ max-width:860px; margin:0 auto; padding:48px 32px; background:#fff; box-shadow:0 1px 3px rgba(0,0,0,0.06); }}
  h1 {{ font-size:1.8em; color:#1a1a2e; margin-bottom:8px; padding-bottom:16px; border-bottom:3px solid #2563eb; }}
  h2 {{ font-size:1.25em; color:#2563eb; margin:36px 0 16px; padding-left:12px; border-left:4px solid #2563eb; }}
  h3 {{ font-size:1.05em; color:#1a1a2e; margin:24px 0 12px; }}
  p {{ color:#475569; margin:8px 0; }}
  blockquote {{ background:#f0f9ff; border-left:4px solid #2563eb; margin:16px 0; padding:12px 20px; border-radius:0 8px 8px 0; color:#475569; }}
  table {{ width:100%; border-collapse:collapse; margin:12px 0; font-size:0.88em; }}
  th {{ background:#eff6ff; color:#2563eb; font-weight:600; padding:10px 14px; text-align:left; border-bottom:2px solid #bfdbfe; }}
  td {{ color:#475569; padding:10px 14px; border-bottom:1px solid #f1f5f9; }}
  tr:hover td {{ background:#f8fafc; }}
  ul, ol {{ margin:8px 0 8px 24px; color:#475569; }}
  li {{ margin:4px 0; }}
  .task-list {{ list-style:none; padding-left:0; }}
  .task-item {{ display:flex; align-items:flex-start; gap:8px; padding:4px 0; }}
  .task-item input[type=checkbox] {{ margin-top:4px; }}
  pre {{ background:#1e293b; color:#e2e8f0; padding:16px 20px; border-radius:8px; overflow-x:auto; font-size:0.85em; line-height:1.6; margin:12px 0; }}
  code {{ font-family:'SF Mono','Fira Code',monospace; background:#f1f5f9; padding:2px 6px; border-radius:4px; font-size:0.9em; }}
  pre code {{ background:none; padding:0; }}
  hr {{ border:none; border-top:1px solid #e2e8f0; margin:24px 0; }}
  .table-wrapper {{ overflow-x:auto; }}
  @media print {{
    body {{ background:#fff; }}
    .guide-container {{ box-shadow:none; padding:0; }}
    h2 {{ break-before:page; }}
  }}
</style>
</head>
<body>
<div class="guide-container">
{body_html}
<p style="text-align:center;color:#94a3b8;font-size:0.78em;margin-top:40px;border-top:1px solid #e2e8f0;padding-top:20px;">
  AI 旅行攻略生成器 · 攻略编号 {guide_id} · 价格信息仅供参考
</p>
</div>
</body>
</html>"""
