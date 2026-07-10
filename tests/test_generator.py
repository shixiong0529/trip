"""
generator.py 覆盖测试：Markdown 解析（含任务一新增的两级嵌套列表）与
HTML/DOCX 输出的关键路径。不触网、不依赖数据库。
"""

import io

import pytest
from docx import Document

from generator import (
    _parse_markdown_to_blocks,
    _blocks_to_html_fragment,
    _extract_stats,
    TravelGuideGenerator,
)


@pytest.fixture
def gen():
    return TravelGuideGenerator("templates")


# ---------- Markdown 解析 ----------

def test_task_list_parses_to_tasks_block():
    md = "- [ ] 预约景点\n- [x] 订酒店\n"
    blocks = _parse_markdown_to_blocks(md)
    assert len(blocks) == 1
    assert blocks[0]["type"] == "tasks"
    assert blocks[0]["items"] == [
        {"checked": False, "text": "预约景点"},
        {"checked": True, "text": "订酒店"},
    ]


def test_non_day_h3_renders_as_sub_title():
    md = "### 预约清单\n- [ ] 景点门票\n"
    blocks = _parse_markdown_to_blocks(md)
    html = _blocks_to_html_fragment(blocks)
    assert '<h3 class="sub-title">预约清单</h3>' in html


def test_script_tag_is_escaped():
    md = "普通段落里混入 <script>alert(1)</script> 恶意内容"
    blocks = _parse_markdown_to_blocks(md)
    html = _blocks_to_html_fragment(blocks)
    assert "<script>" not in html
    assert "&lt;script&gt;" in html


def test_uneven_table_to_docx_does_not_crash(gen):
    md = (
        "# 🗺️ 测试行程\n\n"
        "| 时段 | 安排 |\n"
        "|------|------|\n"
        "| 08:00 | 早餐 | 多出的一列 |\n"
        "| 09:00 | 出发 |\n"
    )
    docx_bytes = gen.to_docx(md, "t1")
    assert len(docx_bytes) > 0
    # 能被 python-docx 重新打开即视为未损坏
    doc = Document(io.BytesIO(docx_bytes))
    tables = doc.tables
    assert len(tables) == 1
    assert len(tables[0].columns) == 3  # 按最宽行（3列）建表，缺的补空


def test_adjacent_highlight_budget_lines_become_two_badges():
    md = (
        "### Day 1 · 测试主题\n"
        "| 时段 | 安排 |\n"
        "|------|------|\n"
        "| 08:00 | 出发 |\n"
        "\n"
        "🎯 **本日亮点：** 参观博物馆\n"
        "💰 **本日预算：** 2人合计 ¥500\n"
    )
    blocks = _parse_markdown_to_blocks(md)
    html = _blocks_to_html_fragment(blocks)
    assert html.count('class="badge badge-green"') == 1
    assert html.count('class="badge badge-orange"') == 1
    assert "参观博物馆" in html
    assert "本日预算" in html


def test_heading_with_own_emoji_does_not_get_extra_icon():
    md = "## 🌤️ 天气与穿搭\n段落内容\n"
    blocks = _parse_markdown_to_blocks(md)
    html = _blocks_to_html_fragment(blocks)
    assert 'class="icon"' not in html
    assert "🌤️ 天气与穿搭" in html


def test_heading_without_emoji_gets_icon():
    md = "## 天气与穿搭\n段落内容\n"
    blocks = _parse_markdown_to_blocks(md)
    html = _blocks_to_html_fragment(blocks)
    assert 'class="icon"' in html
    assert "🌤️" in html


def test_extract_stats():
    text = "3天2人，人均预算 ¥1,850，核心景点 6 个"
    stats = _extract_stats(text)
    stats_dict = dict(stats)
    assert stats_dict["旅行天数"] == "3"
    assert stats_dict["出行人数"] == "2人"
    assert stats_dict["人均预算"] == "¥1,850"
    assert stats_dict["核心景点"] == "6个"


def test_h1_followed_by_plain_paragraph_is_preserved():
    md = "# 🗺️ 目的地行程\n\n这是一段没有统计数字的普通说明文字。\n"
    blocks = _parse_markdown_to_blocks(md)
    html = _blocks_to_html_fragment(blocks)
    assert "这是一段没有统计数字的普通说明文字" in html


def test_intro_stats_are_not_rendered_as_report_cards(gen):
    md = (
        "# 🗺️ 北京4日游\n\n"
        "4天1人，人均预算 ¥2,550，核心景点 8 个\n\n"
        "## 交通建议\n"
        "| 方向 | 推荐方式 |\n"
        "|------|----------|\n"
        "| 去程 | 高铁 |\n"
    )
    html = gen.to_html(md, "no-stats")
    assert 'class="stats-grid"' not in html
    assert 'class="stat-card"' not in html
    assert "旅行天数" not in html
    assert "出行人数" not in html
    assert "人均预算" not in html


def test_report_table_labels_do_not_break_across_lines(gen):
    html = gen.to_html(
        "# 🗺️ 测试\n\n"
        "## 住宿建议\n"
        "| 档位 | 推荐区域/酒店 |\n"
        "|------|---------------|\n"
        "| 经济型 | 测试酒店 |\n",
        "nowrap",
    )
    assert "word-break: keep-all;" in html
    assert "white-space: nowrap;" in html
    assert "min-width: 88px;" in html


def test_nested_ul_parses_to_children():
    md = (
        "- **分层穿搭方案**：\n"
        "  - **上装**：速干T恤\n"
        "  - **下装**：冲锋裤\n"
        "- 简单条目（无子项）\n"
    )
    blocks = _parse_markdown_to_blocks(md)
    assert len(blocks) == 1
    ul = blocks[0]
    assert ul["type"] == "ul"
    assert len(ul["items"]) == 2
    first, second = ul["items"]
    assert first["text"] == "**分层穿搭方案**："
    assert first["children"] == ["**上装**：速干T恤", "**下装**：冲锋裤"]
    assert second["text"] == "简单条目（无子项）"
    assert second["children"] == []


def test_nested_ul_renders_nested_html():
    md = (
        "- **分层穿搭方案**：\n"
        "  - **上装**：速干T恤\n"
        "  - **下装**：冲锋裤\n"
    )
    blocks = _parse_markdown_to_blocks(md)
    html = _blocks_to_html_fragment(blocks)
    assert "<li><strong>分层穿搭方案</strong>：<ul><li><strong>上装</strong>：速干T恤</li>" \
        "<li><strong>下装</strong>：冲锋裤</li></ul></li>" in html


def test_nested_ul_renders_in_docx_with_list_bullet_2(gen):
    md = (
        "# 🗺️ 测试\n\n"
        "- **分层穿搭方案**：\n"
        "  - **上装**：速干T恤\n"
        "  - **下装**：冲锋裤\n"
    )
    docx_bytes = gen.to_docx(md, "t2")
    doc = Document(io.BytesIO(docx_bytes))
    style_names = [p.style.name for p in doc.paragraphs]
    assert "List Bullet" in style_names
    assert "List Bullet 2" in style_names
    # 子项内容确实被写入了对应段落
    child_paras = [p.text for p in doc.paragraphs if p.style.name == "List Bullet 2"]
    assert any("上装" in t for t in child_paras)
    assert any("下装" in t for t in child_paras)


def test_to_html_has_hero_and_balanced_divs(gen):
    md = (
        "# 🗺️ 北京3日游 · 为2人定制\n\n"
        "3天2人，人均预算 ¥1,850，核心景点 6 个\n\n"
        "## 🚄 城际交通建议\n"
        "| 方向 | 推荐方式 |\n"
        "|------|---------|\n"
        "| 去程 | 高铁 |\n\n"
        "### Day 1 · 抵达\n"
        "| 时段 | 安排 |\n"
        "|------|------|\n"
        "| 08:00 | 出发 |\n\n"
        "🎯 **本日亮点：** 到达\n"
        "💰 **本日预算：** ¥100\n"
    )
    html = gen.to_html(md, "t3")
    assert '<div class="hero">' in html
    assert '<div class="container">' in html
    assert html.count("<div") == html.count("</div")


def test_to_html_does_not_insert_route_map(gen):
    md = (
        "# 🗺️ 成都3日游\n\n"
        "3天2人，人均预算 ¥2,000\n\n"
        "**路线总览** 上海 → 成都。\n\n"
        "## 🗺️ 全程路线图\n"
        "按行程顺序绘制的线路示意图。\n\n"
        "```text\n"
        "上海 -> 成都\n"
        "```\n\n"
        "## 🚄 城际交通建议\n"
        "交通内容\n"
    )

    html = gen.to_html(md, "no-route-map")

    assert 'class="route-map-card"' not in html
    assert "全程路线图" not in html
    assert "按行程顺序绘制" not in html
    assert "上海 -&gt; 成都" not in html
    assert "城际交通建议" in html


def test_two_column_tables_use_compact_key_value_layout(gen):
    md = (
        "# 行前物品清单\n\n"
        "## 🎒 行前物品清单\n"
        "| 类别 | 物品 |\n"
        "|------|------|\n"
        "| 🍜 食物 | 自热米饭×6、压缩饼干×1 包、矿泉水 1 箱 |\n"
    )

    html = gen.to_html(md, "kv-table")

    assert 'class="table-wrapper kv-table-wrapper"' in html


def test_overview_route_and_important_tip_use_emphasis_classes():
    md = (
        "# 🗺️ 成都3日游\n\n"
        "3天2人，人均预算 ¥2,000\n\n"
        "**路线总览** 上海 → 成都 → 上海。\n\n"
        "**重要提示：** 本行程为轻松城市漫步，午后避暑。\n"
    )
    blocks = _parse_markdown_to_blocks(md)
    html = _blocks_to_html_fragment(blocks)

    assert 'class="route-overview-card"' in html
    assert 'class="important-note-card"' in html
    assert "上海 → 成都 → 上海" in html


def test_route_overview_card_removes_road_nodes():
    md = (
        "# 🗺️ 川藏行程\n\n"
        "12天2人，人均预算 ¥8,000\n\n"
        "**路线总览：** 成都 → G318川藏南线 → 拉萨 → G109青藏公路 → 西宁。\n"
    )

    html = _blocks_to_html_fragment(_parse_markdown_to_blocks(md))

    assert "成都 → 拉萨 → 西宁" in html
    assert "G318" not in html
    assert "G109" not in html


def test_disclaimer_truncates_followup_content(gen):
    html = gen.to_html(
        "# 🗺️ 成都3日游\n\n"
        "## 行程知识图谱\n"
        "> **免责声明**：请以官方实时信息为准。\n\n"
        "- 是否需要调整某天的景点？\n",
        "truncate-after-disclaimer",
    )

    assert "免责声明" in html
    assert "是否需要调整" not in html
    assert "AI 旅行攻略生成器" not in html


def test_to_html_does_not_embed_homepage_logo(gen):
    html = gen.to_html("# 🗺️ 成都3日游\n\n普通说明\n", "logo")

    assert 'class="brand-logo"' not in html
    assert "data:image/png;base64," not in html
